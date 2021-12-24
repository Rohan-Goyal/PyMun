#! /usr/bin/env python
# Basically the file where I put everything that messes with ODF files
import shutil
import zipfile
from json import load
from math import floor
from os import path
from pathlib import Path
from pprint import pprint
from re import findall
from string import ascii_lowercase
from types import GeneratorType
from urllib.parse import urlparse

import docx
import requests
from bs4 import BeautifulSoup
from docx2python import docx2python
from docx2txt import process as asTxt
from send2trash import send2trash
from urlextract import URLExtract

appname = "pyMUN"


def asArr(filename):
    return docx2python(filename).body


def toPath(
    fileID, app=appname
):  # Given a google drive ID as a string, it returns the "odt" file from which to read, or to which it should be written.
    return f"{path.expanduser('~')}/tmp/{app}/{fileID}.odt"


# We have the document as plaintext. What heuristics can we use to search for it?
"""
Obviously, pull out lists. Either use docx2python or docx to yank out lists. If we have numbered lists, especially if we have sublists, then it's automatically a resolution. No questions.
Position papers never have numbered lists. So if we see no numlists, we have either a note or a position paper. Sources (including official UN resolutions, are PDFS or HTMLs or MHTMLs). DONE: Implement this today.

We can't use wordlimits as a heuristic. So instead, we can use something relating to structuring. Links are a potential heuristic.
We can also consider the number of lines: A bunch of short sentences strongly imply notes/research, and a few discrete paragraphs implies a position paper.
How can we quantify this? Number of line-breaks, perhaps. Or simply use docx to get paragraph counts, and experiment with those. We can also look at length of lines/paras relative to overall doc length.

If a document is composed mostly of numbered lists (especially a single deeply nested list) we can call it a resolution. So basically, we want to create a tree from an un-indented list of paragraphs (as returned by python-docx) based on the resolution structure. If the document matches that structure, call it a resolution.
"""

# def list_elems(doc):
#        paras=[i for i in doc.paragraphs]
#        return [i for i in paras if "list" in i.style.name.lower()]
# TODO: How to get list prefixes
# TODO: Reliable way to get list indentation


def getBody(docArr):
    """Given a nested docarr, skips through the pointless wrapping lists

    :param docArr: Nested array, as returned by docx2python
    :returns: A list of indented/nested doc paragraphs
    :rtype: List

    """
    if len(docArr) > 1:
        return docArr
    else:
        return getBody(docArr[0])


def listElems(docBody):
    """Given a document body, return all lines/paras which are elements of a list

    :param docBody: The document array to search through
    :returns: A list of lines all of which are list elements in a document
    :rtype: List of strings

    """
    # Any element where the first char after indentation is a number, the first two chars are bracketed letter, or a bracketed roman numeral.
    return [
        line
        for line in docBody
        if line.strip() and (line[1] == ")" or hasSmallRoman(line))
    ]  # TODO: Consider numbers


def isResolution(docArr, threshold=0.5):
    """Identify whether a document is a resolution, based on how much of it is wrapped up in lists of some kind.

    :param docArr: The result of passing the target doc through docx2python
    :param threshold: Best not to tinker with it
    :returns: Whether or not it thinks the doc is a resolution, based on the amount of lists in the doc.
    :rtype: Boolean

    """
    # If a document is more than threshold % list, it's a resolution. For now, its 0.5 ie 50%
    return bool(len(listElems(docArr)) >= floor(threshold * len(getBody(docArr))))


def hasSmallRoman(string):
    """Given a string, identify whether it starts with a small roman numeral (i, ii, iv, vii, xi, etc.)

    :param string: The string to check
    :returns: Whether or not that string starts with a small roman numeral
    :rtype: Boolean

    """
    prefix = string.split(")")[0]
    if prefix != prefix.lower():
        return False
    for i in prefix:
        if i not in "ivxl":
            return False
    return True


# For most documents, it has a style "list paragraph" for lists, and 'normal' for the normal ones

# NOTE: Here, the number of \ts before the line begins gives us how deeply indented it is. So what we need is essentially to count the number of \ts before substance, and take the max of that. That gives us how deeply our document is nested. Position papers should be 0, notes can be 0-2, and clauses are 2 or more.


def countLists(docText):
    """Returns the estimated number of top-level lists in a single document

    :param docText: The plaintext of the doc, as return by docx2txt
    :returns: The number of top-level lists
    :rtype: Integer

    """
    # Uses the result of docx2txt rather than 2python:
    # Find the first 1. or 1). When you find one, check the next line.

    # FIXME
    return len(findall("^\s*1\.", docText)) + len(findall("^\s*1\)", docText))


# Find the number of top-level numbered/bulleted lists in a document. Ideally, independent of whether the indentation works or not. So whitespace followed by 1.


def indentLevel(line):
    """Returns the depth of indentation of a given line

    :param line: String representing a single line/para from a word document
    :returns: The level of indentation of the line, -1 if there is no indentation
    :rtype: Integer

    """
    for i, c in enumerate(line):
        if c != "\t":
            return i
    return -1


def maxIndent(docArray):
    """Returns the level of indentation of the most deeply-indented line in a word document

    :param docArray: As returned by getBody
    :returns: The max level of indentation in the document
    :rtype: Integer

    """
    return max([indentLevel(i) for i in docArray])


# The number returned indicates the largest number of 'subs'. So if it returns 2, it means we reach the level of sub-subclauses, i.e sub*2-clauses. Naturally, this only applies if we're actually in a list.
# Hypothesis: Tabs are continuous, i.e a lot at a time. So we basically start from the top, and find the first element that is not a tab. The index of that is what we want
# It works. The problem is that it can't yet tell a top-level list from a top-level paragraph. We could do this, but it's difficult, and depends on how users format their lists.
# NOTE: Doesn't work very well. When I create a doc in Loffice, it works. But for pre-created docs it's not ideal. So for now, keep it as it is.


def links(txt):
    """Return a list of all URLS in a string

    :param txt: Plaintext of a document
    :returns: List of all URLS present
    :rtype: List of string

    """
    return URLExtract().find_urls(txt)


def countLinks(txt):
    return len(links(txt))


def flatten(iterable):
    """Flatten an arbitrarily-nested list or tuple

    :param iterable: The iterable list/tuple to flatten
    :returns: A generator representing the flattened iterable
    :rtype: GeneratorType object

    """
    # WARNING: Does not work on dicts
    it = iter(iterable)
    for e in it:
        if isinstance(e, (list, tuple)):
            for f in flatten(e):
                yield f
        else:
            yield e


def cleanString(string, keyword):
    """Given a string of the form key-value or key:value, cleans it up and obtains the 'value' component of the string. Also truncates it to 75 characters if required.

    :param string: The line in which the key-value construct is present
    :param keyword: The 'key' to search for, such as committee, topic, etc.
    :returns: A string, representing the discovered value (possibly empty)
    :rtype: Str

    """
    # Specifically, cleans a string of the form "Key:value or thereabouts"
    x = string.replace(":", "").replace("-", "").replace(keyword, "").lstrip().rstrip()
    return x if len(x) <= 75 else x[0:74] + "..."


def extractMetadata(docArr):
    """Search through the document array, and attempt to discover document metadata such as the agenda, committee, country.

    :param docArr: An array representing the text of a word document
    :returns: A dictionary of all the metadata it was able to discover.
    :rtype: Dict

    """
    # Search through a flattened version of the list, where one line represents one element.
    # Position papers usually have metadata as key:value somewhere at the top, or occasionally the bottom.
    # https://www.wisemee.com/how-to-write-a-mun-position-paper/
    # https://bestdelegate.com/model-un-made-easy-how-to-write-a-resolution/
    clean = flatten(docArr)  # Each elem of clean is a line of text
    lowered = [i.lower() for i in clean]
    meta = [i for i in lowered if ":" in i or "-" in i]
    agenda = [i for i in meta if "topic" in i]
    committee = [i for i in meta if "committee" in i]
    country = [i for i in meta if "country" in i]
    result = {}
    # Everything up till here is good. The problem now is that resolution headers have committee names as the first line in bold, without any key:value syntax. How to get that
    # Everything below here is hacky and I want to fix it.
    if agenda:
        result.update({"agenda": cleanString(agenda[0], "topic")})
    if committee:
        result.update({"committee": cleanString(committee[0], "committee")})
    if country:
        result.update({"country": cleanString(country[0], "country")})
    return result


def getCommittee(resolution):
    """In resolutions, the committee name is bold and on one of the first lines. This function finds and returns it.

    :param resolution: Path to a document which is an MUN resolution or draft resolution
    :returns: Best guess at committee name
    :rtype: Str

    """
    # Resolutions format the committee name in bold, slightly differently. Searches for it, and returns either a string or None.
    document = docx.Document(resolution)
    bolds = []
    for para in document.paragraphs:
        for run in para.runs:
            if run.bold:
                bolds.append(run.text)
    return bolds[0].replace("Committee:", "").strip() if bolds else None


# Not very robust, but it'll have to do for now.


def wordCount(document):
    """Count the number of words in a word document

    :param document: An array, as returned by getBody.
    :returns: The number of words in the document, including lists, etc.
    :rtype: Integer

    """
    # Document is an array body as returned by getBody(docArr.body)
    body = list(flatten(document))
    body = " ".join(body).split()
    return len(body)


def docType(docxFile):
    """Uses a variety of heuristics to deduce the type of a given document, returning unclassified if it fails

    :param docxFile: Path to the word document that needs to be classified
    :returns: One of ("resolution","position","notes","unclassified")
    :rtype: String

    """
    docArr = getBody(asArr(docxFile))
    maxdepth = maxIndent(docArr)  # NOTE: Unreliable
    if maxdepth >= 2 or bool(
        len(listElems(docArr)) >= floor(0.5 * len(docArr))  # Reliable
    ):  # More than 50% list
        return "resolution"
    for i in docArr:
        if hasSmallRoman(i):  # We have a 3-nested sublist
            return "resolution"
    # We can develop the stuff above a bit, but for now it's fine
    meta = extractMetadata(docArr)
    if "country" in meta:
        return "position"
    if countLinks(docxFile) > 1:
        return "notes"
    # TODO: Para structure tests

    if wordCount(docArr) >= 900:
        return "notes"

    return "unclassified"


def magicParse(path):
    """Determine the type of a document, and extract whatever metadata it can

    :param path: Document to parse/analyse
    :returns: Dictionary of metadata it has been able to infer
    :rtype: Dict

    """
    documentType = docType(path)
    metadata = extractMetadata(getBody(asArr(path)))
    return (
        {
            "committee": getCommittee(path),
            **metadata,
            "type": documentType,
        }
        if documentType == "resolution" and getCommittee(path)
        else {**metadata, "type": documentType}
    )


def customClassify(title, localPath):
    """Read the config.json file and find the list of custom rules. Then, apply them to a given file

    :param title: The title of the file, as it is stored in Gdrive
    :param localPath: The path to the download of the file.
    :returns: A classification for the document if the custom rules provide one, else None
    :rtype: String, or Nonetype

    """
    # Read json. Get a list of custom rules and split them into two dicts: containRules, nameRules
    with open("config.json") as conf:
        rules = load(conf)["custom-rules"]
    nameRules = {i["regex"]: i["type"] for i in rules["name"]}
    containRules = {i["regex"]: i["type"] for i in rules["contains"]}
    for k, v in nameRules.items():
        if k in title:
            return v
    text = asTxt(localPath)
    for k, v in containRules.items():
        if k in text:
            return v
    return None  # Or false


# Credit https://github.com/python-openxml/python-docx/issues/610
def getHtmlData(html):
    """Given a well-formed HTML page, return metadata including the title, and whatever other metadata the document contains

    :param html: The html to scrape/extract metadata from
    :returns: Dictionary containing whatever metadata was extracted
    :rtype: Dict

    """
    # Given a URL, return a dict of the title, site/source, and any other metadata
    soup = BeautifulSoup(html, features="lxml")
    title = soup.title.string
    metas = soup.find_all("meta")
    metadata = {}
    relevant = [
        m for m in metas if "property" in m.attrs and "og" in m.attrs["property"]
    ]
    for meta in relevant:
        metadata.update(
            {meta.attrs["property"].replace("og:", ""): meta.attrs["content"]}
        )
    return {"web_title": title, **metadata}


def getLinkData(url):
    """Given a URL, extract whatever metadata is possible, including the title, source, and whatever metadata the webpage itself provides

    :param url: URL to read
    :returns: Key-value representation of website metadata attributes
    :rtype: Dict

    """
    try:
        response = requests.get(url, timeout=2)
    except:
        response = False
    source = urlparse(url).netloc.replace(
        "www.", ""
    )  # Consider replacing the last bit of the domain name (.org, etc)
    if response:
        return {**getHtmlData(response.content), "source": source}
    else:
        return {"source": source, "web_title": url.split("/")[-1]}


# DONE: Test on some HTML pages (actual sources I used for MUN)
"""
Plan for XML parsing:
Extract the zips to a temp file. Mess with the ./word/document.xml using a very simple find and replace system.
Compress the folder back and write it to the original docx file
DONE"""


def linkDict(txt):
    """Given a string, find all links in that string, and create string representations of them of the form "{title} {source}".

    :param txt: The string/text from which to extract links
    :returns: Dict, where keys are links found in the text and values are f-strings containing some metadata
    :rtype: Dict

    """
    allLinks = [i for i in links(txt) if "schemas" not in i and "http" in i]
    # print(allLinks)
    conversionTable = {}
    for link in allLinks:
        meta = getLinkData(link)
        string = f"\"{meta['web_title']} ({meta['source']}) [{link}]\""  # consider replacing webtitle
        conversionTable.update({link: string})
    return conversionTable


def extractToTemp(docPath, appName="pyMUN"):
    """Extract a docx file to a folder, and return the folder path

    :param docPath: The word doc to unzip
    :param appName: Decides the subfolder of ~/tmp to look in when saving the folder
    :returns: Path to the extracted folder
    :rtype: Path

    """
    z = zipfile.ZipFile(docPath)
    zipPath = Path(
        f"{path.expanduser('~')}/tmp/{appName}/zips/{z.filename.split('/')[-1].replace('.docx','')}"
    )  # ~/tmp/pymun/zips/{docID} is the format of the path
    z.extractall(zipPath)
    return zipPath


def writeToDoc(folderPath):
    """Write a folder (the result of unzipping a word doc) to a word doc

    :param folderPath: Path of the unzipped word doc
    :returns: Path to the created docx file
    :rtype: Path

    """
    docPath = folderPath.resolve().parents[1] / (folderPath.name + ".docx")
    # pyMUN folder, then find the filename
    d = shutil.make_archive(docPath, "zip", folderPath)
    shutil.move(d, docPath)
    return docPath


def getDocumentFile(folderPath):
    """Given the path to an unzipped word doc, find actual document.xml file

    :param folderPath:Path to word doc (zip file)
    :returns: Path to the document.xml file contained within it
    :rtype: String (path)

    """
    return (
        f"{folderPath}word/document.xml"
        if str(folderPath)[-1] == "/"
        else f"{folderPath}/word/document.xml"
    )


def replaceLinksXml(filePath):
    """Given an XML file (document.xml, specifically), replace all the hyperlinks within it

    :param filePath: Path to the document.xml
    :returns: A string of the XML file, with all links replaced
    :rtype: String

    """
    # Enough libraries, I'll just mess with the XML directly and rewrite it. Time to do something on the level of parsing HTML with regex.
    with open(filePath, "r") as xmlFile:
        string = xmlFile.read()
        lines = xmlFile.readlines()
    links = linkDict(string)
    newString = string
    for i in links:
        newString = newString.replace(i, links[i])
    return newString


def replaceLinks(docPath):
    """Given a path to a document, uses XML chicanery to replace the displayed text of the various links to something which reflects link metadata (as returned by linkDict() func)

    :param docPath: Path to the docx file which needs to be manipulated
    :returns: None
    :rtype: NoneType

    """
    folder = extractToTemp(docPath)
    f = getDocumentFile(folder)
    replaced = replaceLinksXml(f)
    with open(f, "w+") as xmlFile:
        xmlFile.write(replaced)
    writeToDoc(folder)
    # DONE: Somehow mark this file as one that needs to be re-uploaded to google drive. basically, my instinct is to somehow wrap it as a google drive file object, and then push that object (or a reference to it) to a file or a list. So push the path to a file or list
    send2trash(str(folder))


class Tree:
    # Attributes: A list of children, all of whom are either trees or empty.
    def __init__(self, rootstring, parent=None):
        self.root = rootstring
        self.children = []
        self.parent = parent

    def __str__(self):
        return self.root  # TODO: Consider fixing

    def addChild(self, child):
        self.children.append(Tree(child, self))

    def addChildren(self, children):
        self.children.extend([Tree(i, self) for i in children])

    def removeChild(self, child):
        if child in self.children:
            self.children.remove(child)

    def maxDepth(self):
        if self.children:
            depths = [i.maxDepth() for i in self.children]
            return max(depths) + 1
        else:
            return 0
        # Tree with only one node (no children) registers as depth 0.

    def getChild(self, n):
        return self.children[n]

    def getNestedChild(self, *args):
        if len(args) == 1:
            return self.getChild(args[0])
        else:
            return self.getChild(args[0]).getNestedChild(*args[1:])

    def flatten(
        self,
    ):  # Return a list of all clauses and subclauses in order. No formatting, no nesting, just all the strings.
        yield self.root
        for i in self.children:
            yield i.flatten()
            # for j in i.children:
            #        yield j.flatten()

    def flattenGenerators(
        self, flat
    ):  # The flatten function leaves nested lists as generators
        for i, v in enumerate(flat):
            if isinstance(v, GeneratorType):
                flat[i] = self.flattenGenerators(list(v))
        return flatten(flat)

    def fullFlatten(self):
        return self.flattenGenerators(list(self.flatten()))


def treeTest():
    x = Tree("Hello")
    x.addChildren(["0", "1", "2"])
    x.getChild(0).addChild("3")
    x.getNestedChild(0, 0).addChild("000")
    x.getNestedChild(0, 0).addChild("0000")
    print(x)
    pprint(x.fullFlatten())


class Clause(Tree):
    @staticmethod
    def filter_clauses(body, testfunc):
        indices = [i for i, v in enumerate(body) if testfunc(v)]
        values = [Tree(body[i]) for i in indices]
        return indices, values

    @staticmethod
    def fromFormattedDocArr(
        docArr,
    ):  # For now. Assumes everything is perfectly formatted
        body = getBody(docArr)
        docstart = Tree("Document Start")
        toplevel = listElems(body)
        headers = [
            i for i in body if i not in clauses
        ]  # In theory, just the preamble stuff. In practice, a lot more.

        indices, clauses = Clause.filter_clauses(
            clauses, lambda i: i[0] in "123456789" and i[1] in ")."
        )
        for index, root in enumerate(clauses):
            # Everything between the current toplevel and the next toplevel
            children = toplevel[
                indices[index] : indices[index + 1]
            ]  # So between the first and second clauses, or between the second and third,or whatever
            subindices, subclauses = Clause.filter_clauses(
                children,
                lambda i: i[0] == "(" and i[2] == ")" and i[1] in ascii_lowercase,
            )

            for subindex, subroot in enumerate(subclauses):
                subchildren = children[subindices[subindex] : subindices[subindex + 1]]
                subsubindices, subsubclauses = Clause.filter_clauses(
                    subchildren, hasSmallRoman
                )  # NOTE: Since MUN only allows up to sub-sub-sub-clauses, we're probably fine without a filter and just taking everything. But keep this, partly to test if it works and partly because... reasons.
                subclauses[subindex].addChildren(subsubclauses)

            clauses[index].addChildren(subclauses)
        docstart.addChildren(clauses)
        return docstart

    # TODO: This is a reliable-ish way to get a tree, at the cost of decent formatting and the like.
    # For simplicity, and because recursion, we'll consider the document a single clause with subclauses
    # We've seen that docx2python makes it a nested list and uses tabs for indents. So we'll need to start at the top,
    # NOTE: We can consider using
    @staticmethod
    def appendOrReplace(target, replacer):
        # Check if the last char is punctuation. If so, replace it with replacer. Otherwise, append replacer to target.
        last = target[-1]
        if last in "!&*-;:,.?":
            target[-1] = replacer
            return target
        else:
            return target + replacer

    def format(self):
        # Regardless of nesting level, we basically have a semicolon at the end of every sub* clause except the last one, which has a.
        # So for every clause, I'd suggest children[-1] have a fullstop added. For the rest, if they have children, append a ":", else append a ";"
        # If the last char is a punctuation mark, replace it, otherwise append.
        flat = self.fullFlatten()
        body = self.root
        if flat[-1] == body:
            # Last clause
            self.appendOrReplace(".")
        if self.children:
            self.appendOrReplace(":")
        else:
            self.appendOrReplace(";")

        for i in self.children:
            i.format()
            # Calls itself recursively. Stops when there are no children left, which is perfect


# TODO: Write a tree to a word doc somehow.

# NOTE: Should consider using the 1) a) i) system to find lists, as well as simply indentation. Perhaps combine both and come up with a nice heuristic.
# NOTE: It seems periods are only for the very last element.

"""Get the document as a tree structure: So headings, subheadings, lists, sublists"""
# TODO: Clause objects, autoformatting
